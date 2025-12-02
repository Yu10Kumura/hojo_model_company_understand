"""
各種フォーマットへのエクスポート

改良点:
 - PDF: fpdf(旧)からfpdf2のUnicode対応を前提とした実装へ拡張
 - 日本語フォント(TTF)が存在する場合は全文字を保持
 - フォント未配置時はASCIIフォールバックし、Streamlit側で警告表示できるようメッセージ返却を想定
"""
import json
from datetime import datetime
from docx import Document
from fpdf import FPDF
import re
from pathlib import Path


JAPANESE_FONT_CANDIDATES = [
    # プロジェクト内配置想定
    Path(__file__).parent.parent / "fonts" / "NotoSansJP-Regular.ttf",
    Path(__file__).parent.parent / "fonts" / "IPAexGothic.ttf",
    # ユーザが手動追加するかもしれない場所の例 (macOS Home ディレクトリ配下など)
    Path.home() / "Library" / "Fonts" / "NotoSansJP-Regular.ttf",
    Path.home() / "Library" / "Fonts" / "IPAexGothic.ttf",
]


def _find_japanese_font() -> Path | None:
    for p in JAPANESE_FONT_CANDIDATES:
        if p.is_file():
            return p
    return None


def export_to_json(
    company_name: str,
    final_report: str
) -> str:
    """
    JSON形式でエクスポート
    
    Args:
        company_name: 会社名
        final_report: 完全版レポート
    
    Returns:
        JSON文字列
    """
    data = {
        "company_name": company_name,
        "generated_at": datetime.now().isoformat(),
        "report": final_report
    }
    
    return json.dumps(data, ensure_ascii=False, indent=2)


def export_to_word(company_name: str, content: str):
    """
    分析結果をWord文書として出力
    """
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    
    doc = Document()
    
    # タイトル
    title = doc.add_heading(f'{company_name} 企業分析レポート', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 現在の日付
    date_para = doc.add_paragraph(f'作成日: {datetime.now().strftime("%Y年%m月%d日")}')
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # コンテンツを行ごとに処理
    lines = content.split('\n')
    in_table = False
    table = None
    
    for line in lines:
        line = line.strip()
        
        # 表の開始を検出（統合アウトプット表）
        if line.startswith('| カテゴリ') and '分析項目' in line:
            in_table = True
            # 5列のテーブルを作成
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # ヘッダー行を設定
            header_cells = table.rows[0].cells
            headers = ['カテゴリ', '分析項目', 'ファクト/推論', '採用への示唆', '確認要否']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # ヘッダーを太字に
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            continue
            
        # 表の区切り行をスキップ
        elif line.startswith('| :---') or line.startswith('|:---'):
            continue
            
        # 表のデータ行を処理
        elif in_table and line.startswith('|') and line.endswith('|'):
            # データを分割
            cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
            
            if len(cells_data) == 5:  # 5列のデータがある場合
                # 新しい行を追加
                row = table.add_row()
                for i, cell_data in enumerate(cells_data):
                    # Markdownの太字記号を除去
                    clean_data = cell_data.replace('**', '')
                    row.cells[i].text = clean_data
            continue
            
        # 表の終了を検出
        elif in_table and (line == '' or line.startswith('#')):
            in_table = False
            table = None
            
        # 通常のテキスト処理（表以外）
        if not in_table:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line:
                doc.add_paragraph(line)
    
    return doc


def export_to_pdf(
    company_name: str,
    final_report: str
) -> bytes:
    """
    PDF形式でエクスポート(日本語対応版)

    改善ポイント:
      - 日本語フォント(TTF)検出 → Unicode対応フォントとして登録
      - 未検出時はASCIIフォールバック (日本語は "?" に置換)
      - 呼び出し側でフォント未検出を警告表示可能にするため、レポート末尾に簡易注記を付与

    フォント設置手順例:
      fonts/ ディレクトリを作成し、以下いずれかを配置してください。
        - NotoSansJP-Regular.ttf (推奨: Google Fonts)
        - IPAexGothic.ttf (IPA フォント)

    Args:
        company_name: 会社名
        final_report: 完全版レポート

    Returns:
        PDFバイナリ
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()

    font_path = _find_japanese_font()
    japanese_enabled = False
    if font_path:
        try:
            # 名称 "JP" で追加
            pdf.add_font("JP", "", str(font_path), uni=True)
            pdf.set_font("JP", size=14)
            japanese_enabled = True
        except Exception:
            # フォント追加失敗時はフォールバック
            pdf.set_font("Arial", size=14)
    else:
        pdf.set_font("Arial", size=14)

    # タイトル
    title = "採用戦略分析レポート" if japanese_enabled else "Recruitment Strategy Report"
    pdf.set_font(pdf.font_family, size=18)
    pdf.cell(0, 10, txt=title, ln=True, align='C')
    pdf.set_font(pdf.font_family, size=11)
    company_label = f"対象企業: {company_name}" if japanese_enabled else f"Company: {company_name}"
    generated_label = (
        f"生成日時: {datetime.now().strftime('%Y-%m-%d %H:%M')}" if japanese_enabled
        else f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    pdf.cell(0, 8, txt=company_label, ln=True)
    pdf.cell(0, 8, txt=generated_label, ln=True)
    pdf.ln(4)

    # 本文フォントを少し小さく
    pdf.set_font(pdf.font_family, size=10)

    for raw_line in final_report.split("\n"):
        line = raw_line.rstrip()
        if not line:
            pdf.ln(4)
            continue

        # 見出しレベル検出 (Markdown記法)
        heading_level = 0
        if line.startswith("### "):
            heading_level = 3
            line = line[4:].strip()
        elif line.startswith("## "):
            heading_level = 2
            line = line[3:].strip()
        elif line.startswith("# "):
            heading_level = 1
            line = line[2:].strip()

        if heading_level:
            size_map = {1: 14, 2: 12, 3: 11}
            pdf.set_font(pdf.font_family, size=size_map.get(heading_level, 12))
            pdf.multi_cell(0, 7, txt=line)
            pdf.set_font(pdf.font_family, size=10)
            continue

        # テーブル行やその他複雑記法は簡易的にスキップ/整形
        if line.startswith("|") and line.endswith("|"):
            continue

        if not japanese_enabled:
            # 日本語未対応フォントの場合はASCII以外を '?' に置換
            safe_line = ''.join(ch if ord(ch) < 128 else '?' for ch in line)
        else:
            safe_line = line

        try:
            pdf.multi_cell(0, 5, txt=safe_line)
        except Exception:
            # 念のためフォールバック
            fallback = ''.join(ch if ord(ch) < 128 else '?' for ch in safe_line)
            pdf.multi_cell(0, 5, txt=fallback)

    if not japanese_enabled:
        pdf.ln(6)
        pdf.set_font(pdf.font_family, size=8)
        pdf.multi_cell(0, 4, txt="[Japanese font not found: Non-ASCII characters were replaced with '?'. Place a TTF font in fonts/ to enable full Japanese support.]")

    # 出力
    # fpdf2では output() がバイナリを返す(Python3)
    try:
        return pdf.output()
    except Exception as e:
        # フォールバック: エラー時は簡易版PDF生成
        error_msg = str(e)[:100]
        print("[PDF ERROR] " + error_msg + ", creating ASCII fallback")
        pdf_fallback = FPDF()
        pdf_fallback.add_page()
        pdf_fallback.set_font("Arial", size=12)
        pdf_fallback.cell(0, 10, txt="Report Generation Error", ln=True)
        pdf_fallback.multi_cell(0, 6, txt="Japanese font is required. Please install a TrueType font in fonts/ directory.")
        return pdf_fallback.output()
